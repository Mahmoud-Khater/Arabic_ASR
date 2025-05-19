import time
import os
import whisper
import torch
from django.shortcuts import render
# from pyannote.audio import Pipeline
from pydub import AudioSegment
import uuid
from django.core.files.storage import FileSystemStorage
from django.http import StreamingHttpResponse
from docx import Document
import json
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, AutoModel
from transformers import AutoProcessor, AutoModelForCTC, pipeline,WhisperProcessor,WhisperForConditionalGeneration,WhisperFeatureExtractor,WhisperTokenizer
import torchaudio
# from multiple_datasets.hub_default_utils import convert_hf_whisper

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print("Device:", device)

#torch.backends.cuda.matmul.allow_tf32 = True  # Enable TF32 for matrix mult
#torch.backends.cudnn.allow_tf32 = True        # Enable TF32 for cuDNN ops
# إعدادات الجهاز (الـ GPU أو الـ CPU)

# model_path = "D:/STT_OCR_RAG/audio_transcription/TranslationModel" #win
model_path = "/media/nozom/New Volume1/STT_OCR_RAG/audio_transcription/TranslationModel" #linux

print("Loading Trans model...")
model_startTime = time.time()
model_trans = AutoModelForSeq2SeqLM.from_pretrained(model_path).to(device)
tokenizer_trans = AutoTokenizer.from_pretrained(model_path)
model_trans.eval()
model_endTime = time.time()
elapsed_model_time = model_endTime - model_startTime
print(f"⏱️ وقت المودل: {elapsed_model_time:.2f} ثانية")
# local_model_path_fine = "D:/STT_OCR_RAG/audio_transcription/voice_to_text/results_small/checkpoint-100" #win
local_model_path_fine = "/media/nozom/New Volume1/STT_OCR_RAG/audio_transcription/results_small/checkpoint-900" #linux

start = time.time()
# model = WhisperForConditionalGeneration.from_pretrained(local_model_path_fine)
# feature_extractor = WhisperFeatureExtractor.from_pretrained(local_model_path_fine)
# processor = WhisperProcessor.from_pretrained(local_model_path_fine, language="ar", task="transcribe")
# model.generation_config.forced_decoder_ids = None
# tokenizer = WhisperTokenizer.from_pretrained(local_model_path_fine, language="ar", task="transcribe")
# model.generation_config.task = "transcribe"
model = whisper.load_model(
"medium",
device=device,
download_root="./models"
)
model.to(device)
end = time.time()
print("Model small time:", end - start)

print("Loading Turbo model...")
start = time.time()
model_cache = whisper.load_model(
"turbo",
device=device,
download_root="./models"
)
end = time.time()
print("Model turbo loading time:", end - start)


if not torch.cuda.is_available():
    raise RuntimeError("CUDA GPU is not available. Please check your PyTorch installation.")

supported_langs = {
    "en": "eng_Latn",
    "fr": "fra_Latn",
    "de": "deu_Latn",
    "es": "spa_Latn",
    "zh": "zho_Hans",
    "ru": "rus_Cyrl",
    "ar": "arb_Arab",
    "tr": "tur_Latn",
    "it": "ita_Latn",
    "ja": "jpn_Jpan",
    "he": "heb_Hebr"
}
def clean_transcript(text):
    bad_phrases = [
        "اشتركوا في القناة",
        "لا تنسوا الإعجاب",
        "فعلوا زر الجرس",
        "رابط في الوصف",
        "ترجمة نانسى",
        "Merci d'avoir regardé cette vidéo",
        "C'est parti !"
    ]
    if any(phrase in text for phrase in bad_phrases):
        return True
    return False

def remove_consecutive_duplicates(text):
    words = text.split()
    result = []

    for i, word in enumerate(words):
        if i == 0 or word != words[i - 1]:
            result.append(word)

    return ' '.join(result)
def transcribe_long_audio(m,lang,count,audio_path, num, ar, start, end, output_txt):
    if not os.path.exists(audio_path):
        raise FileNotFoundError(f"{audio_path} does not exist!")

    # استخدام FP16 إذا كان متاح
    fp16 = torch.cuda.is_available()
    if ar:
        lang = 'ar'
    else:
        if count == 0 and m==0:
            lang = None
        else:
            lang = lang
    flag = 0
    if ar:
        result = model_cache.transcribe(
        audio_path,
        fp16=fp16,
        language="ar",
        verbose=True
        )
#         result = pipeline(
#     "automatic-speech-recognition",
#     model=model,
#     tokenizer=processor.tokenizer,
#     feature_extractor=processor.feature_extractor,
#     device="cuda" if torch.cuda.is_available() else "cpu"
# )(audio_path, generate_kwargs={
#         "language": "ar",
#         "task": "transcribe",
#         "return_timestamps": True  # Enable timestamps for long audio
#     })
        f = clean_transcript(result["text"])
        result["text"] = remove_consecutive_duplicates(result["text"])
        if f:
            result["text"] = "------------------"
            translation = "------------------"
    else:
        flag=1
        result = model_cache.transcribe(
            audio_path,
            fp16=fp16,
            language=lang,
            verbose=True
        )
        f = clean_transcript(result["text"])
        if f:
            result["text"] = "------------------"
            translation = "------------------"
        if not f:
            start_time = time.time()
            tokenizer_trans.src_lang = supported_langs[result['language']]
            forced_bos_token_id = tokenizer_trans.convert_tokens_to_ids("arb_Arab")

            with torch.no_grad():
                inputs = tokenizer_trans(result["text"], return_tensors="pt", truncation=True).to(device)
                inputs["forced_bos_token_id"] = forced_bos_token_id
                output = model_trans.generate(**inputs, max_new_tokens=256)
                translation = tokenizer_trans.batch_decode(output, skip_special_tokens=True)[0]
                del inputs, output
                torch.cuda.empty_cache()
            elapsed_time = time.time() - start_time
            print(f"⏱️ وقت الترجمة: {elapsed_time:.2f} ثانية")
    if ar:
        lang_to_print = "ar"
    else:
        lang_to_print = result['language']
    text = result["text"]
    data = []
    translations = []

    if output_txt:
        word_file = "TR.docx"
        try:
            doc = Document(word_file)
        except Exception as e:
            print("Error opening document:", e)
            doc = Document()

        if flag == 1 :

            doc.add_paragraph(text)
            doc.add_paragraph(translation)

            data.append(text)
            translations.append(translation)
        else:
            doc.add_paragraph(text)
            data.append(text)
        doc.save(word_file)

    return data, lang_to_print , num , translations

def format_time(seconds):
    minutes = int(seconds // 60)
    sec = seconds % 60
    return f"{minutes:02}:{int(sec)}"  # e.g., 01:05

def process_audio(files,request):
    for i,audio_file in enumerate(files):
        m=0
        file_name_org = audio_file.name
        language = request.POST.get(f'language_{i}', 'any')
        if language == 'arabic' :
            ar = 1
            lang = "ar"
        else :
            ar=0
            lang = language
            m=1

        fs = FileSystemStorage(location='uploads')
        filename = f"{uuid.uuid4()}.wav"
        filepath = os.path.join('uploads', filename)
        fs.save(filename, audio_file)
        print(filepath)
        audio = AudioSegment.from_file(filepath)
        new_data = []
        num = 0
        speaker_segments = []
        duration_ms = len(audio)
        print(f"Full recording duration: {duration_ms / 1000} seconds")
        start_ms = 0
        chunk_size = 30000
        overlap = 0
        end_ms = chunk_size
        while start_ms < duration_ms:
            real_end_ms = min(end_ms, duration_ms)
            segment = audio[start_ms:real_end_ms]
            speaker_segments.append({
                "speaker": "all_speakers",
                "start": start_ms / 1000,
                "end": real_end_ms / 1000,
                "audio": segment
            })
            start_ms = end_ms - overlap
            end_ms = start_ms + chunk_size
        count = 0
        for i, segment in enumerate(speaker_segments, start=1):
            temp_path = f"temp_{i}_{file_name_org}.wav"
            try:
                # وقت تصدير الصوت
                segment["audio"].export(temp_path, format="wav")
                start_s = format_time(segment["start"])
                end_s = format_time(segment["end"])
                transcribe_start_time = time.time()
                if count == 0 and language=="any":
                    lang = None
                    m=1
                torch.cuda.empty_cache()
                text, lan, num , trans= transcribe_long_audio(m,lang,count,temp_path, num, ar, start_s, end_s, output_txt=True)
                count = count +1
                lang = lan
                transcribe_time = time.time() - transcribe_start_time
                print(f"Transcription time for segment {i}: {transcribe_time:.2f} seconds")

                if i == 1:
                    new_data.append(f"Language is : {lan}")
                    yield json.dumps({"new_item": new_data[-1], "type": "language"}) + "\n"
                    new_data.append(f"Sound is : {file_name_org}")
                    yield json.dumps({"new_item": new_data[-1], "type": "address"}) + "\n"

                for item in text:
                    new_data.append({
                        "time": f"{start_s} -> {end_s}",
                        "speaker": item
                    })
                    yield json.dumps({
                        "new_item": new_data[-1],
                        "type": "time_speaker"
                    }) + "\n"
                for item in trans:
                    new_data.append({
                        "time": f"{start_s} -> {end_s}",
                        "speaker": item
                    })
                    yield json.dumps({
                        "new_item": new_data[-1],
                        "type": "time_trans"
                    }) + "\n"

            finally:
                # تأكد من حذف الملف المؤقت بعد الانتهاء
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    print(f"Deleted temporary file: {temp_path}")
        # os.remove(temp_path)
        os.remove(filepath)
        lang = None
def index(request):
    # model = load_whisper_model_gpu()
    # tran = load_trans_model_gpu()
    start = time.time()
    if request.method == 'GET':
        return render(request, 'index.html')

    if request.method == 'POST' and request.FILES.getlist('audio_file'):
        audio_files  = request.FILES.getlist('audio_file')
        return StreamingHttpResponse(process_audio(audio_files,request), content_type='application/json')
    end = time.time()
    print("Total time*******************---------->>:", end - start)
    return render(request, 'index.html')
